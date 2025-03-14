import sys
import os
import csv
from datetime import datetime
import re
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.tag import pos_tag
import pdfplumber
from PyPDF2 import PdfReader
from docx import Document
from pathlib import Path
from fastapi import APIRouter, HTTPException, Body
from fastapi.responses import JSONResponse
from config import UPLOAD_DIR, EXTRACTION_DIR, DATASET_DIR, OUTPUT_DIR, LOG_DIR
import logging
from typing import List
from pydantic import BaseModel
from utils.file_utils import get_file_security_classification
import itertools
import json
from nltk.tokenize.punkt import PunktSentenceTokenizer
from nltk.chunk import ne_chunk
from nltk.corpus import names

# Get the current file's directory
current_dir = Path(__file__).parent
nltk_data_dir = current_dir / 'nltk_data'

# Configure NLTK data path
nltk.data.path.append(str(nltk_data_dir))

# Configure logging
logging.basicConfig(filename=LOG_DIR / 'extraction_service.log', level=logging.INFO)
logger = logging.getLogger(__name__)

def has_dependent_pronoun(sentence):
    """Check if a sentence starts with a pronoun or contains dependent markers."""
    words = word_tokenize(sentence.lower())
    if not words:
        return False
        
    # Common pronouns that might indicate dependency
    pronouns = {'he', 'she', 'it', 'they', 'his', 'her', 'their', 'its', 'this', 'that', 'these', 'those'}
    
    # Words/phrases that might indicate a dependent clause
    dependent_markers = {'however', 'therefore', 'thus', 'hence', 'consequently', 'as a result',
                        'moreover', 'furthermore', 'in addition', 'additionally', 'also',
                        'meanwhile', 'nevertheless', 'nonetheless', 'instead', 'conversely',
                        'rather', 'although', 'though', 'otherwise', 'likewise', 'similarly',
                        'accordingly', 'subsequently'}
    
    # Check first word for pronouns
    if words[0] in pronouns:
        return True
        
    # Check for dependent markers at start of sentence
    first_phrase = ' '.join(words[:3]).lower()
    if any(marker in first_phrase for marker in dependent_markers):
        return True
        
    # Check for other indicators of dependency
    tagged = pos_tag(words)
    if tagged and tagged[0][1] in {'PRP', 'PRP$', 'DT'}:  # Personal pronouns, possessive pronouns, or determiners
        return True
        
    return False

def extract_sentences(text):
    """
    Extract sentences while preserving context for dependent clauses.
    """
    try:
        # Clean the text first
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        text = text.strip()
        
        # Use NLTK's sentence tokenizer
        sentences = sent_tokenize(text)
        if not sentences:
            return []
            
        # Process sentences to combine dependent ones with their context
        processed_sentences = []
        i = 0
        while i < len(sentences):
            current_sentence = sentences[i].strip()
            combined_sentence = current_sentence
            
            # Look ahead for dependent sentences
            while i + 1 < len(sentences) and has_dependent_pronoun(sentences[i + 1]):
                # Combine with next sentence
                combined_sentence += " " + sentences[i + 1].strip()
                i += 1
            
            if combined_sentence:
                processed_sentences.append(combined_sentence)
            i += 1
        
        return processed_sentences
        
    except Exception as e:
        logger.error(f"Error in sentence extraction: {str(e)}")
        return [text] if text.strip() else []

def extract_sentences_with_context(paragraph):
    """
    Extract sentences from a paragraph while maintaining the portion marking context.
    Returns a tuple of (portion_marking, list_of_sentences)
    """
    try:
        # First, extract any portion marking from the paragraph
        portion_marking, clean_paragraph = extract_portion_marking(paragraph)
        
        # Now extract sentences from the cleaned paragraph
        sentences = extract_sentences(clean_paragraph)
        
        return portion_marking, sentences
    except Exception as e:
        logger.error(f"Error in sentence extraction with context: {str(e)}")
        return None, [paragraph] if paragraph.strip() else []

router = APIRouter()

@router.get("/files/")
async def list_files(folder: str = ""):
    try:
        target_dir = UPLOAD_DIR / folder if folder else UPLOAD_DIR
        logger.info(f"Checking directory: {target_dir}")
        
        if not target_dir.exists():
            logger.error(f"Directory not found: {target_dir}")
            raise HTTPException(status_code=404, detail=f"Directory not found: {target_dir}")
            
        files = []
        logger.info(f"Scanning directory contents...")
        
        # First, build a map of metadata information
        metadata_map = {}
        for metadata_file in target_dir.glob('*.metadata'):
            try:
                with open(metadata_file, 'r') as f:
                    metadata_json = json.loads(f.read().strip())
                    base_filename = metadata_file.stem.replace('.metadata', '')
                    # Extract just the security classification
                    metadata_map[base_filename] = metadata_json.get('security_classification', 'UNCLASSIFIED')
            except Exception as e:
                logger.error(f"Error reading metadata file {metadata_file}: {str(e)}")
                metadata_map[base_filename] = 'UNCLASSIFIED'
        
        for item in target_dir.iterdir():
            # Skip hidden files and metadata files
            if item.name.startswith('.') or item.name.endswith('.metadata'):
                logger.info(f"Skipping file: {item.name}")
                continue
                
            # Only include supported file types for extraction
            if item.is_file() and item.suffix.lower() not in ['.pdf', '.docx', '.txt']:
                logger.info(f"Skipping unsupported file: {item.name}")
                continue
                
            logger.info(f"Adding file to list: {item.name}")
            
            # Get classification from metadata
            base_filename = item.stem
            classification = metadata_map.get(base_filename, "UNCLASSIFIED")
            
            file_info = {
                "name": item.name,
                "path": item.name,  # Just use filename since all files are in uploads directory
                "type": "folder" if item.is_dir() else item.suffix[1:].lower(),
                "size": item.stat().st_size if item.is_file() else None,
                "classification": classification
            }
            
            if item.is_file():
                file_info["uploadDate"] = datetime.fromtimestamp(item.stat().st_mtime).isoformat()
                
            files.append(file_info)
            
        # Sort files by upload date (newest first)
        files.sort(key=lambda x: x.get("uploadDate", ""), reverse=True)
        
        logger.info(f"Returning {len(files)} files")
        logger.info(f"File list: {files}")
        
        return files
        
    except Exception as e:
        logger.error(f"Error listing files: {str(e)}")
        logger.exception("Full traceback:")
        raise HTTPException(status_code=500, detail=str(e))

class ExtractionRequest(BaseModel):
    filenames: List[str]
    csv_filename: str

def extract_text_with_layout(file_path):
    """Extract text while preserving layout information."""
    _, ext = os.path.splitext(file_path)
    logger.info(f"=== Starting text extraction for file: {file_path} ===")
    logger.info(f"File extension: {ext.lower()}")
    
    try:
        if ext.lower() == '.pdf':
            content = extract_from_pdf(file_path)
        elif ext.lower() == '.docx':
            content = extract_from_docx(file_path)
        elif ext.lower() == '.txt':
            content = extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
            
        logger.info(f"Extraction completed. Content sections: {len(content)}")
        logger.info(f"First section preview: {content[0][:200] if content else 'No content'}")
        return content
    except Exception as e:
        logger.error(f"Error in text extraction: {str(e)}")
        logger.exception("Full traceback:")
        raise

def extract_from_pdf(file_path):
    logger.info(f"=== Extracting from PDF: {file_path} ===")
    extracted_content = []
    try:
        # First attempt: Using PDFPlumber which often preserves layout better
        with pdfplumber.open(file_path) as pdf:
            logger.info(f"PDF opened successfully with pdfplumber. Total pages: {len(pdf.pages)}")
            
            # Extract each page with layout preservation
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text with better whitespace handling
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                
                if text:
                    # Process page breaks and potential paragraph splits
                    if extracted_content and not extracted_content[-1].endswith(('.', '!', '?', ':', ';')):
                        # This might be a continuation from the previous page
                        # Check if current page starts with lowercase letter or without proper capitalization
                        if text and (not text[0].isupper() or text[0] in ',:;)]}'):
                            # Append to the previous content without page break
                            extracted_content[-1] += " " + text
                            logger.info(f"Page {page_num}: Joined with previous page content")
                            continue
                    
                    logger.info(f"Page {page_num}: Extracted {len(text)} characters")
                    extracted_content.append(text)
                else:
                    logger.warning(f"Page {page_num}: No text extracted with pdfplumber")
        
        # If pdfplumber couldn't extract anything useful, try PyPDF4 as fallback
        if not any(content.strip() for content in extracted_content):
            logger.info("No content extracted with pdfplumber, trying PyPDF4 as fallback")
            extracted_content = []
            
            with open(file_path, 'rb') as file:
                reader = PdfFileReader(file)
                for page in range(reader.getNumPages()):
                    text = reader.getPage(page).extractText()
                    if text and text.strip():
                        # Normalize newlines and spaces
                        text = re.sub(r'\s+', ' ', text)
                        extracted_content.append(text.strip())
        
        # Post-process content to clean up and ensure paragraph detection will work
        processed_content = []
        for content in extracted_content:
            # Replace multiple spaces with single space
            content = re.sub(r' +', ' ', content)
            # Add paragraph breaks for bullet points and numbered lists
            content = re.sub(r'([.!?])\s+([•\-*]|\d+\.|\([a-z]\))', r'\1\n\n\2', content)
            # Add paragraph breaks for clear paragraph transitions
            content = re.sub(r'([.!?])\s+([A-Z][a-z]+)', r'\1\n\n\2', content)
            processed_content.append(content)
        
        if not processed_content:
            logger.warning("No content extracted from PDF with any method")
            return ["[PDF text extraction failed. This PDF might be scanned or have content that's not easily extractable.]"]
            
        logger.info(f"PDF extraction complete. Total sections: {len(processed_content)}")
        return processed_content
    except Exception as e:
        logger.error(f"Error extracting PDF: {str(e)}")
        logger.exception("Full traceback:")
        raise

def extract_from_docx(file_path):
    logger.info(f"=== Extracting from DOCX: {file_path} ===")
    extracted_content = []
    current_section = ""
    
    try:
        doc = Document(file_path)
        logger.info(f"DOCX opened successfully. Total paragraphs: {len(doc.paragraphs)}")
        
        # Keep track of heading levels and sections
        current_heading_level = 0
        in_list = False
        list_level = 0
        
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                # Empty paragraph might indicate section break
                if current_section:
                    extracted_content.append(current_section)
                    current_section = ""
                continue
            
            # Check if this is a heading
            is_heading = False
            if para.style.name.startswith('Heading'):
                is_heading = True
                try:
                    level = int(para.style.name.replace('Heading', ''))
                    current_heading_level = level
                except ValueError:
                    current_heading_level = 1  # Default to level 1 if cannot parse
            
            # Check if this is a list item
            is_list_item = False
            if para.style.name.startswith(('List', 'Bullet', 'Numbered')):
                is_list_item = True
                # Try to determine list level
                if para.paragraph_format.left_indent:
                    list_level = int(para.paragraph_format.left_indent / 360)  # Approximate
                else:
                    list_level = 0
            
            # Start a new section if:
            # 1. This is a heading
            # 2. This is a bullet point/list item
            # 3. Previous paragraph ended with sentence-ending punctuation
            if (is_heading or 
                is_list_item or 
                (current_section and current_section.rstrip().endswith(('.', '!', '?')))
               ):
                if current_section:
                    extracted_content.append(current_section)
                    current_section = ""
            
            # Add the paragraph to the current section
            if current_section:
                # Add a double newline for headings and list items to help with paragraph detection
                if is_heading or is_list_item:
                    current_section += "\n\n" + text
                else:
                    current_section += " " + text
            else:
                current_section = text
        
        # Add the last section if there's any content
        if current_section:
            extracted_content.append(current_section)
        
        # Ensure we've extracted content
        if not extracted_content:
            logger.warning("No content extracted from DOCX")
            extracted_content = ["[DOCX file appears to be empty or has no readable text content]"]
        
        logger.info(f"DOCX extraction complete. Total sections: {len(extracted_content)}")
        return extracted_content
    except Exception as e:
        logger.error(f"Error extracting DOCX: {str(e)}")
        logger.exception("Full traceback:")
        raise

def extract_from_txt(file_path):
    logger.info(f"=== Extracting from TXT: {file_path} ===")
    try:
        # Attempt to detect encoding
        encodings = ['utf-8', 'latin-1', 'windows-1252', 'ascii']
        content = None
        
        # Try different encodings
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                logger.info(f"TXT file read successfully with {encoding} encoding. Total characters: {len(content)}")
                break
            except UnicodeDecodeError:
                logger.warning(f"Failed to read with {encoding} encoding, trying next...")
                continue
        
        if not content:
            logger.error("Could not read text file with any supported encoding")
            return ["[Error: Could not determine the text encoding of this file]"]
        
        # Pre-process text to detect document structure
        # Replace Windows line endings with Unix line endings
        content = content.replace('\r\n', '\n')
        
        # Look for section separators (common in many text documents)
        section_separators = ['---', '***', '===', '___', '#####']
        
        # Split by section separators if they are present
        for separator in section_separators:
            if separator in content:
                logger.info(f"Found section separator: {separator}")
                # Replace separator with a form feed character to mark section boundaries
                content = re.sub(r'\s*' + re.escape(separator) + r'\s*', '\f', content)
        
        # Split into sections
        sections = content.split('\f')
        
        # If we only have one section, add more contextual splits based on line breaks
        if len(sections) == 1:
            processed_content = content
            
            # Add paragraph breaks before lines that look like headers
            processed_content = re.sub(r'\n([A-Z][A-Z\s]+:?)\n', r'\n\n\1\n', processed_content)
            
            # Add paragraph breaks after numbered sections
            processed_content = re.sub(r'(\d+\..*?)\n', r'\1\n\n', processed_content)
            
            # Preserve existing paragraph breaks (multiple newlines)
            processed_content = re.sub(r'\n{3,}', '\n\n', processed_content)
            
            return [processed_content]
        else:
            # Process each section
            processed_sections = []
            for section in sections:
                if section.strip():
                    processed_sections.append(section.strip())
            
            logger.info(f"TXT extraction split into {len(processed_sections)} sections")
            return processed_sections
            
    except Exception as e:
        logger.error(f"Error extracting TXT: {str(e)}")
        logger.exception("Full traceback:")
        raise

def extract_paragraphs(text):
    """
    Enhanced paragraph extraction that handles various document formats and potential issues.
    Uses multiple heuristics to identify paragraph boundaries even in poorly formatted documents.
    """
    # Check if the text is empty or None
    if not text or not text.strip():
        return []
    
    # First attempt: Split by double newlines (original method)
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    
    # If we got reasonable paragraphs, return them
    if len(paragraphs) > 1 and all(len(p.split()) > 3 for p in paragraphs):
        return paragraphs
    
    # Second attempt: Try single newlines, but be smarter about it
    # This helps when documents use single line breaks between paragraphs
    candidate_paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    
    # Merge lines that are likely part of the same paragraph
    merged_paragraphs = []
    current_paragraph = ""
    
    for i, line in enumerate(candidate_paragraphs):
        # Start a new paragraph if:
        # 1. This is the first line
        # 2. Previous line ends with a period, question mark, or exclamation point
        # 3. Current line starts with a capital letter and previous is non-empty
        # 4. Line is very short (likely a header)
        # 5. Line has special formatting (bullet points, numbers, etc.)
        
        starts_new_paragraph = (
            i == 0 or 
            (i > 0 and candidate_paragraphs[i-1] and 
             any(candidate_paragraphs[i-1].rstrip().endswith(c) for c in ['.', '!', '?', ':', ';'])) or
            (i > 0 and line and line[0].isupper() and len(line.split()) > 3) or
            (len(line.split()) < 3 and len(line) < 25) or
            line.lstrip().startswith(('•', '-', '*', '1.', '2.', '3.', 'a.', 'b.', 'c.'))
        )
        
        if starts_new_paragraph and current_paragraph:
            merged_paragraphs.append(current_paragraph)
            current_paragraph = line
        else:
            # Add a space if we're appending to existing content
            if current_paragraph:
                current_paragraph += " " + line
            else:
                current_paragraph = line
    
    # Add the last paragraph
    if current_paragraph:
        merged_paragraphs.append(current_paragraph)
    
    # If we got reasonable paragraphs, return them
    if len(merged_paragraphs) > 1:
        return merged_paragraphs
    
    # Third attempt: Use sentence structure for paragraph detection
    # This is a fallback when document has no clear paragraph formatting
    sentences = extract_sentences(text)
    
    # Group sentences into paragraphs
    if len(sentences) <= 1:
        return [text.strip()] if text.strip() else []
    
    # Use topic shifts to identify paragraph boundaries
    # Simplified approach: look for shifts in content or structure
    paragraphs = []
    current_paragraph = sentences[0]
    
    for i in range(1, len(sentences)):
        current_sentence = sentences[i]
        previous_sentence = sentences[i-1]
        
        # Check for potential paragraph break indicators
        new_paragraph = (
            # Significant change in sentence length (could indicate a style shift)
            abs(len(current_sentence.split()) - len(previous_sentence.split())) > 10 or
            
            # Current sentence starts with transition words or phrases
            any(current_sentence.lower().startswith(phrase) for phrase in 
                ["however", "furthermore", "meanwhile", "in conclusion", "finally",
                 "therefore", "thus", "consequently", "in summary", "in contrast"]) or
            
            # Every 3-5 sentences (approximate approach for when no other signals exist)
            i % 4 == 0
        )
        
        if new_paragraph:
            paragraphs.append(current_paragraph)
            current_paragraph = current_sentence
        else:
            current_paragraph += " " + current_sentence
    
    # Add the last paragraph
    if current_paragraph:
        paragraphs.append(current_paragraph)
    
    return paragraphs or [text.strip()] if text.strip() else []

def is_meaningful_text(text):
    """
    Enhanced check for meaningful text that considers context and completeness.
    """
    words = word_tokenize(text)
    tagged = pos_tag(words)
    
    # Check for minimum requirements
    has_noun = any(tag.startswith('NN') for _, tag in tagged)
    has_verb = any(tag.startswith('VB') for _, tag in tagged)
    is_long_enough = len(words) > 5
    
    # Calculate text quality metrics
    text_chars = sum(1 for c in text if c.isalpha())
    total_chars = len(text)
    text_ratio = text_chars / total_chars if total_chars > 0 else 0
    
    # Check sentence structure
    has_subject_predicate = False
    for i, (word, tag) in enumerate(tagged):
        if tag.startswith('NN'):  # Found a noun
            # Look for a verb after the noun
            for _, vtag in tagged[i+1:]:
                if vtag.startswith('VB'):
                    has_subject_predicate = True
                    break
            if has_subject_predicate:
                break
    
    return (has_noun and has_verb and is_long_enough and 
            text_ratio > 0.7 and has_subject_predicate)

def clean_text(text):
    """
    Advanced text cleaning that removes headers, footers, page numbers, 
    and other artifacts that might interfere with paragraph detection.
    """
    # Remove common header/footer patterns
    cleaned_text = remove_headers_footers(text)
    
    # Remove headers, footers, and page numbers
    cleaned_text = re.sub(r'^.*?Page \d+.*?$', '', cleaned_text, flags=re.MULTILINE)
    cleaned_text = re.sub(r'^.*?\d+/\d+.*?$', '', cleaned_text, flags=re.MULTILINE)  # Page X of Y format
    
    # Remove excessive whitespace
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
    
    # Remove watermark text that might appear
    common_watermarks = ['confidential', 'draft', 'internal use only', 'do not copy', 'not for distribution']
    for watermark in common_watermarks:
        cleaned_text = re.sub(r'(?i)' + watermark, '', cleaned_text)
    
    return cleaned_text.strip()

def remove_headers_footers(text):
    """
    Detect and remove repeating headers and footers from document text.
    Uses pattern detection to identify repeated lines at the top or bottom of pages.
    """
    # Split text into lines
    lines = text.split('\n')
    if len(lines) <= 5:  # Not enough lines to detect patterns
        return text
    
    # Find potential headers (repeated at the beginning of text sections)
    potential_headers = []
    sections = text.split('\f')  # Form feed character often separates pages
    if len(sections) > 1:
        # Get first line from each section
        section_starts = [s.split('\n')[0].strip() for s in sections if s.strip()]
        
        # Count occurrences of each starting line
        from collections import Counter
        start_counter = Counter(section_starts)
        
        # Lines that appear as the first line in multiple sections are likely headers
        potential_headers = [line for line, count in start_counter.items() 
                            if count > 1 and len(line) > 5]
    
    # Find potential footers (repeated at the end of text sections)
    potential_footers = []
    if len(sections) > 1:
        # Get last line from each section
        section_ends = [s.split('\n')[-1].strip() for s in sections if s.strip()]
        
        # Count occurrences of each ending line
        end_counter = Counter(section_ends)
        
        # Lines that appear as the last line in multiple sections are likely footers
        potential_footers = [line for line, count in end_counter.items() 
                            if count > 1 and len(line) > 5]
    
    # Remove identified headers and footers
    cleaned_lines = []
    for line in lines:
        line_stripped = line.strip()
        # Skip if line matches a header or footer
        if any(header == line_stripped for header in potential_headers) or \
           any(footer == line_stripped for footer in potential_footers):
            continue
        # Skip lines that are just page numbers
        if re.match(r'^\s*\d+\s*$', line_stripped) or \
           re.match(r'^\s*Page\s+\d+\s*$', line_stripped, re.IGNORECASE) or \
           re.match(r'^\s*\d+\s+of\s+\d+\s*$', line_stripped, re.IGNORECASE):
            continue
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def create_csv_metadata(csv_filename: str, original_name: str, source_files: List[str], source_classifications: List[str]):
    """Create metadata file for a CSV file."""
    metadata = {
        "filename": csv_filename,
        "original_name": original_name,
        "created_at": datetime.now().isoformat(),
        "source_files": [
            {
                "filename": filename,
                "security_classification": classification
            }
            for filename, classification in zip(source_files, source_classifications)
        ]
    }
    
    metadata_path = EXTRACTION_DIR / f"{csv_filename}.metadata"
    with open(metadata_path, 'w') as f:
        json.dump(metadata, f, indent=2)

def get_file_security_classification(filename):
    """Get security classification from file's metadata."""
    try:
        # Handle potential folder paths in filename
        filename = filename.replace('/', os.sep).replace('\\', os.sep)
        
        # Get the full path components
        file_path = Path(filename)
        file_stem = file_path.stem
        folder_path = file_path.parent
        
        # Construct metadata path
        if str(folder_path) != '.':
            metadata_path = UPLOAD_DIR / folder_path / f"{file_stem}.metadata"
        else:
            metadata_path = UPLOAD_DIR / f"{file_stem}.metadata"
            
        logger.info(f"=== Security Classification Lookup ===")
        logger.info(f"Input filename: {filename}")
        logger.info(f"Upload directory: {UPLOAD_DIR}")
        logger.info(f"File stem: {file_stem}")
        logger.info(f"Folder path: {folder_path}")
        logger.info(f"Looking for metadata at: {metadata_path}")
        logger.info(f"Metadata exists: {metadata_path.exists()}")
        
        if metadata_path.exists():
            logger.info(f"Found metadata file at: {metadata_path}")
            with open(metadata_path, 'r', encoding='utf-8') as f:
                try:
                    metadata = json.load(f)
                    logger.info(f"Metadata content: {json.dumps(metadata, indent=2)}")
                    classification = metadata.get('security_classification')
                    if classification:
                        logger.info(f"Found classification: {classification}")
                        return classification
                    else:
                        logger.warning("No security_classification field in metadata")
                except json.JSONDecodeError as e:
                    logger.error(f"Error decoding metadata JSON: {e}")
                    logger.info(f"Raw metadata content: {f.read()}")
                    raise
        
        # If we get here, either no metadata file or no classification found
        logger.warning(f"No valid metadata found, using default classification")
        logger.warning(f"Available metadata files: {list(UPLOAD_DIR.glob('**/*.metadata'))}")
        return 'UNCLASSIFIED'
    except Exception as e:
        logger.error(f"Error reading metadata for {filename}: {str(e)}")
        logger.exception("Full traceback:")
        return 'UNCLASSIFIED'

def extract_portion_marking(text):
    """
    Extract portion marking from the beginning of text if present.
    Portion markings appear between parentheses at the start of paragraphs.
    Normalizes and standardizes common classification markings.
    """
    try:
        # Simple pattern to match any content within parentheses at the start of text
        portion_marking_pattern = r'^\s*\(([^)]+)\)\s*'
        
        match = re.match(portion_marking_pattern, text)
        if match:
            # Extract the classification marking (everything inside the parentheses)
            classification = match.group(1).strip()
            
            # First, convert everything to uppercase for consistent processing
            classification = classification.upper()
            
            # 1. Basic classification level normalization with word boundaries
            if re.match(r'^\s*U\s*$', classification):
                classification = 'UNCLASSIFIED'
            elif re.match(r'^\s*C\s*$', classification):
                classification = 'CONFIDENTIAL'
            elif re.match(r'^\s*S\s*$', classification):
                classification = 'SECRET'
            elif re.match(r'^\s*TS\s*$', classification):
                classification = 'TOP SECRET'
            
            # 2. Standardize separators first (before other replacements)
            classification = re.sub(r'\s*[/]+\s*', '//', classification)
            
            # 3. Replace full forms (with word boundaries)
            classification = re.sub(r'\bFOR\s+OFFICIAL\s+USE\s+ONLY\b', 'FOUO', classification)
            classification = re.sub(r'\bSENSITIVE\s+COMPARTMENTED\s+INFORMATION\b', 'SCI', classification)
            
            # 4. Normalize control markings (with word boundaries)
            classification = re.sub(r'\bNF\b', 'NOFORN', classification)
            classification = re.sub(r'\bREL\s+TO\b', 'REL TO', classification)
            
            # 5. Standardize special access program markings (with word boundaries)
            classification = re.sub(r'\bSI(?:-[A-Z])?\b', 'SI', classification)
            classification = re.sub(r'\bTK\b', 'TK', classification)
            classification = re.sub(r'\bHCS(?:-[A-Z])?(?:-P)?\b', 'HCS', classification)
            
            # 6. Clean up any multiple separators that might have been created
            classification = re.sub(r'[/]{2,}', '//', classification)
            
            # 7. Clean up whitespace
            classification = re.sub(r'\s+', ' ', classification.strip())
            
            # Remove the portion marking from the text
            clean_text = re.sub(portion_marking_pattern, '', text, 1).strip()
            
            logger.info(f"Extracted portion marking: {classification}")
            return classification, clean_text
            
        return None, text
    except Exception as e:
        logger.error(f"Error extracting portion marking: {str(e)}")
        return None, text

@router.post("/extract/")
async def extract_file_content(request_data: ExtractionRequest):
    logger.info("=== Starting new extraction process ===")
    logger.info(f"Files to process: {request_data.filenames}")
    logger.info(f"Output CSV filename: {request_data.csv_filename}")
    
    results = []
    extracted_content = []
    source_classifications = []

    for filename in request_data.filenames:
        try:
            file_path = UPLOAD_DIR / filename.replace('/', os.sep).replace('\\', os.sep)
            logger.info(f"\n=== Processing file: {filename} ===")
            logger.info(f"Full file path: {file_path}")
            logger.info(f"File exists: {file_path.exists()}")
            
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                results.append({"filename": filename, "status": "File not found"})
                continue

            # Get security classification
            file_security_classification = get_file_security_classification(filename)
            source_classifications.append(file_security_classification)
            logger.info(f"File security classification: {file_security_classification}")

            # Extract content
            raw_content = extract_text_with_layout(str(file_path))
            logger.info(f"Raw content extracted: {len(raw_content)} sections")

            # Process paragraphs
            logger.info("=== Processing paragraphs ===")
            paragraph_count = 0
            meaningful_paragraph_count = 0
            
            for content in raw_content:
                cleaned_content = clean_text(content)
                paragraphs = extract_paragraphs(cleaned_content)
                logger.info(f"Found {len(paragraphs)} paragraphs in section")
                
                for para in paragraphs:
                    paragraph_count += 1
                    if is_meaningful_text(para):
                        meaningful_paragraph_count += 1
                        content_classification, clean_para = extract_portion_marking(para)
                        
                        logger.info(f"Processing paragraph {paragraph_count}:")
                        logger.info(f"Original length: {len(para)}")
                        logger.info(f"Cleaned length: {len(clean_para)}")
                        logger.info(f"Classification found: {content_classification}")
                        
                        final_classification = content_classification if content_classification else file_security_classification
                        
                        extracted_content.append({
                            "answer": clean_para if content_classification else para,
                            "source": filename,
                            "file_security_classification": file_security_classification,
                            "content_security_classification": final_classification,
                            "type": "paragraph"
                        })
                    else:
                        logger.debug(f"Paragraph {paragraph_count} filtered out as not meaningful")

            logger.info(f"Paragraph processing complete:")
            logger.info(f"Total paragraphs: {paragraph_count}")
            logger.info(f"Meaningful paragraphs: {meaningful_paragraph_count}")

            # Process sentences
            logger.info("=== Processing sentences ===")
            sentence_count = 0
            meaningful_sentence_count = 0
            
            for content in raw_content:
                cleaned_content = clean_text(content)
                paragraphs = extract_paragraphs(cleaned_content)
                
                for para in paragraphs:
                    if is_meaningful_text(para):
                        para_classification, sentences = extract_sentences_with_context(para)
                        para_final_classification = para_classification if para_classification else file_security_classification
                        
                        logger.info(f"Found {len(sentences)} sentences in paragraph")
                        
                        for sentence in sentences:
                            sentence_count += 1
                            if is_meaningful_text(sentence):
                                meaningful_sentence_count += 1
                                extracted_content.append({
                                    "answer": sentence,
                                    "source": filename,
                                    "file_security_classification": file_security_classification,
                                    "content_security_classification": para_final_classification,
                                    "type": "sentence"
                                })
                            else:
                                logger.debug(f"Sentence {sentence_count} filtered out as not meaningful")

            logger.info(f"Sentence processing complete:")
            logger.info(f"Total sentences: {sentence_count}")
            logger.info(f"Meaningful sentences: {meaningful_sentence_count}")

            results.append({"filename": filename, "status": "Content extracted successfully"})
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            logger.exception("Full traceback:")
            results.append({"filename": filename, "status": f"Error: {str(e)}"})

    # Prepare CSV output
    logger.info("\n=== Preparing CSV Output ===")
    logger.info(f"Total extracted items: {len(extracted_content)}")
    if extracted_content:
        logger.info("Sample of extracted content:")
        for i, item in enumerate(extracted_content[:2]):
            logger.info(f"Item {i + 1}:")
            logger.info(f"  Type: {item['type']}")
            logger.info(f"  Source: {item['source']}")
            logger.info(f"  Answer preview: {item['answer'][:100]}...")
            logger.info(f"  Classification: {item['content_security_classification']}")

    # Create CSV file
    current_date = datetime.now().strftime('%Y%m%d_%H%M%S')
    csv_filename = f"{request_data.csv_filename}_{current_date}.csv"
    csv_path = EXTRACTION_DIR / csv_filename

    try:
        logger.info(f"Writing CSV file: {csv_path}")
        with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
            csv_writer = csv.writer(csvfile)
            csv_writer.writerow([
                "question", 
                "answer", 
                "source", 
                "file security classification",
                "content security classification",
                "type"
            ])
            rows_written = 0
            for item in extracted_content:
                csv_writer.writerow([
                    "",  # question
                    item["answer"],
                    item["source"],
                    item["file_security_classification"],
                    item["content_security_classification"],
                    item["type"]
                ])
                rows_written += 1
            
            logger.info(f"CSV file written successfully. Total rows: {rows_written}")
    except Exception as e:
        logger.error(f"Error writing CSV file: {str(e)}")
        logger.exception("Full traceback:")
        raise

    # Create metadata file
    try:
        logger.info("Creating metadata file")
        create_csv_metadata(
            csv_filename=csv_filename,
            original_name=csv_filename,
            source_files=request_data.filenames,
            source_classifications=source_classifications
        )
        logger.info("Metadata file created successfully")
    except Exception as e:
        logger.error(f"Error creating metadata file: {str(e)}")
        logger.exception("Full traceback:")
        raise

    logger.info("=== Extraction process completed ===")
    return {
        "status": "Extraction process completed",
        "results": results,
        "csv_file": str(csv_path.name),
        "extracted_items_count": len(extracted_content)
    }

@router.get("/csv-files/")
async def get_csv_files():
    try:
        csv_files = []
        for f in os.listdir(EXTRACTION_DIR):
            if f.endswith('.csv'):
                file_path = EXTRACTION_DIR / f
                metadata_path = EXTRACTION_DIR / f"{f}.metadata"
                
                file_info = {
                    "name": f,
                    "created_at": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat()
                }
                
                # Add metadata information if available
                if metadata_path.exists():
                    with open(metadata_path, 'r') as mf:
                        metadata = json.load(mf)
                        file_info["original_name"] = metadata.get("original_name")
                        file_info["source_files"] = metadata.get("source_files")
                
                csv_files.append(file_info)
                
        return JSONResponse(content=csv_files, status_code=200)
    except Exception as e:
        logger.error(f"Error fetching CSV files: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error fetching CSV files: {str(e)}")

@router.get("/csv-preview/{filename}")
async def get_csv_preview(filename: str, rows: int = 1000):  # Default to 1000 rows, but allow it to be configurable
    try:
        file_path = EXTRACTION_DIR / filename
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="CSV file not found")
        
        with open(file_path, 'r', newline='', encoding='utf-8') as csvfile:
            csv_reader = csv.DictReader(csvfile)
            data = list(itertools.islice(csv_reader, rows))  # Use itertools.islice for efficiency
        
        return JSONResponse(content=data, status_code=200)
    except Exception as e:
        logger.error(f"Error fetching CSV preview for {filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error fetching CSV preview: {str(e)}")

@router.post("/rename-csv/")
async def rename_csv_file(request: dict = Body(...)):
    old_name = request.get("old_name")
    new_name = request.get("new_name")
    
    if not old_name or not new_name:
        raise HTTPException(status_code=400, detail="Both old_name and new_name are required")
    
    old_path = EXTRACTION_DIR / old_name
    new_path = EXTRACTION_DIR / new_name
    
    if not old_path.exists():
        raise HTTPException(status_code=404, detail="CSV file not found")
    
    if new_path.exists():
        raise HTTPException(status_code=400, detail="New file name already exists")
    
    try:
        # Rename CSV file
        old_path.rename(new_path)
        
        # Update metadata file
        old_metadata_path = EXTRACTION_DIR / f"{old_name}.metadata"
        new_metadata_path = EXTRACTION_DIR / f"{new_name}.metadata"
        
        if old_metadata_path.exists():
            with open(old_metadata_path, 'r') as f:
                metadata = json.load(f)
            
            # Update filename in metadata
            metadata["filename"] = new_name
            
            # Write updated metadata to new file
            with open(new_metadata_path, 'w') as f:
                json.dump(metadata, f, indent=2)
            
            # Remove old metadata file
            old_metadata_path.unlink()
        
        return JSONResponse(
            content={"message": f"CSV file renamed from '{old_name}' to '{new_name}' successfully"},
            status_code=200
        )
    except Exception as e:
        logger.error(f"Error renaming CSV file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error renaming CSV file: {str(e)}")

@router.delete("/delete-csv/{filename}")
async def delete_csv_file(filename: str):
    file_path = EXTRACTION_DIR / filename
    metadata_path = EXTRACTION_DIR / f"{filename}.metadata"
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="CSV file not found")
    
    try:
        # Delete CSV file
        os.remove(file_path)
        
        # Delete metadata file if it exists
        if metadata_path.exists():
            os.remove(metadata_path)
            
        return JSONResponse(
            content={"message": f"CSV file '{filename}' and its metadata deleted successfully"},
            status_code=200
        )
    except Exception as e:
        logger.error(f"Error deleting CSV file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting CSV file: {str(e)}")

@router.post("/update-csv/{filename}")
async def update_csv_file(filename: str, data: dict = Body(...)):
    try:
        csv_path = EXTRACTION_DIR / filename
        if not csv_path.exists():
            raise HTTPException(status_code=404, detail=f"CSV file '{filename}' not found")

        logger.info(f"Updating CSV file: {filename}")
        logger.info(f"Received data: {json.dumps(data, indent=2)}")

        # First, verify we can read the existing file
        with open(csv_path, 'r', newline='', encoding='utf-8') as csvfile:
            reader = csv.DictReader(csvfile)
            original_fieldnames = reader.fieldnames
            logger.info(f"Original CSV fieldnames: {original_fieldnames}")

        # Now write the updated data
        with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=original_fieldnames)
            writer.writeheader()
            
            if data.get('data'):
                logger.info(f"Writing {len(data['data'])} rows to CSV")
                for row in data['data']:
                    # Map the row data to match the CSV field names exactly
                    csv_row = {}
                    for field in original_fieldnames:
                        csv_row[field] = row.get(field, '')
                    
                    logger.info(f"Writing row: {csv_row}")
                    writer.writerow(csv_row)

        # Verify the file was written correctly
        with open(csv_path, 'r', newline='', encoding='utf-8') as csvfile:
            content = csvfile.read()
            logger.info(f"CSV file content after write:\n{content}")

        return JSONResponse(
            content={"message": f"CSV file '{filename}' updated successfully"},
            status_code=200
        )
    except Exception as e:
        logger.error(f"Error updating CSV file {filename}: {str(e)}")
        logger.exception("Full traceback:")
        raise HTTPException(status_code=500, detail=f"Error updating CSV file: {str(e)}")